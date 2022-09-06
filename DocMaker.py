"""
DocMaker:

Two classes:
    1) BuildCSV creates a CSV file that lists all jpg and png files in a given folder
    2) Scaler scales the images in that CSV file and saves them in a docx document.
"""

from docx import Document
from docx.shared import Cm, Mm
from numpy import nan
import pandas as pd
from os import listdir
from os.path import join
from natsort import natsorted
from PIL import ImageOps, ImageDraw
import tempfile
from ImageOrientation import ImageOrientation


class BuildCSV:
    def __init__(self, image_files):
        self.image_files = image_files
        self.image_files_cleaned = self.clean()

    def clean(self) -> list:
        assert isinstance(self.image_files, list)
        new_img_lst = []
        for img in listdir(self.image_files[0]):
            if img[-4:].lower() in ['.jpg', '.png']:
                new_img_lst.append(join(self.image_files[0], img))
        sorted_list = natsorted(new_img_lst)
        return sorted_list

    def build(self, filename:str, scale:float) -> str:

        dataframe = pd.DataFrame({'Image_file': self.image_files_cleaned})
        dataframe['Scale'] = float(scale)
        dataframe['Artwork_life_size_height'] = nan
        dataframe['Artwork_life_size_width'] = nan
        dataframe['Artwork_frame_height'] = nan
        dataframe['Artwork_frame_width'] = nan
        dataframe['Frame_type'] = 'None'
        dataframe.to_csv(f'{filename}.csv')

        return f'{filename}.csv'


class Scaler:
    def __init__(self, csv_file, save_location, docx_filename):

        #print(csv_file, save_location, docx_filename)
        self.csv_file = csv_file  # location of the csv file
        self.save_location = save_location  # location to save the docx
        self.docx_filename = join(self.save_location, f'{docx_filename}.docx')  # final file name
        self.document = Document()
        section = self.document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        self.temp_dir = tempfile.TemporaryDirectory()
        self.image_orientation = None
        self.error_msg = []

    def iterate_images(self):
        df = pd.read_csv(self.csv_file)

        for index, row in df.iterrows():
            try:
                if row['Image_file']:
                    if str(row['Artwork_frame_height']) == 'nan':
                        self.no_frame(row)
                    else:
                        if row['Frame_type'].lower() in ['none', 'rectangle', 'square']:
                            self.add_rectangle_frame(row)
                        elif row['Frame_type'].lower() == 'circle' or row['Frame_type'].lower() == 'oval':
                            self.add_ellipse_frame(row)

                    self.document.add_paragraph(str(row['Image_file']))
            except ValueError as e:
                self.error_msg.append(f"Image {row['Image_file']} encountered the following error: {e} "
                                      f"- Check heights and widths are correct")
                #print(e)
                continue

            except AttributeError as e:
                self.error_msg.append(f"Image {row['Image_file']} encountered the following error: {e} "
                                      f"- Contact ilkka at ilkkasipila@outlook.com")
                #print(e)
                continue

    def calculate_size_in_cm(self, object_size: float, scale: int) -> float:
        target_size_in_cm = object_size / scale
        return target_size_in_cm

    def run_scaler(self):
        self.iterate_images()
        self.document.save(self.docx_filename)

    def get_error_msgs(self):
        return self.error_msg

    def no_frame(self, row):
        real_height = row['Artwork_life_size_height']
        real_width = row['Artwork_life_size_width']
        scale = row['Scale']
        target_img_height_in_cm = self.calculate_size_in_cm(real_height, scale)
        target_img_width_in_cm = self.calculate_size_in_cm(real_width, scale)
        picture_name = row['Image_file']
        self.image_orientation = ImageOrientation(picture_name)
        picture = self.image_orientation.re_orient()
        picture.save(join(self.temp_dir.name, 'temp_img.jpg'))
        self.document.add_picture(join(self.temp_dir.name, 'temp_img.jpg'), height=Cm(target_img_height_in_cm),
                                  width=Cm(target_img_width_in_cm))

    def add_rectangle_frame(self, row):
        padded_img = self._get_padded_img(row)
        framed_img = ImageOps.expand(padded_img, border=5)

        scale = row['Scale']
        target_img_height_in_cm = self.calculate_size_in_cm(row['Artwork_frame_height'], scale)
        target_img_width_in_cm = self.calculate_size_in_cm(row['Artwork_frame_width'], scale)

        framed_img.save(join(self.temp_dir.name, 'temp_img.jpg'))
        self.document.add_picture(join(self.temp_dir.name, 'temp_img.jpg'), height=Cm(target_img_height_in_cm),
                                  width=Cm(target_img_width_in_cm))

    def add_ellipse_frame(self, row):
        padded_img = self._get_padded_img(row)
        framed_img = ImageDraw.Draw(padded_img)
        padded_img_width, padded_img_height = padded_img.size
        framed_img.ellipse([0, 0, padded_img_width, padded_img_height], outline='black', width=5)
        scale = row['Scale']
        target_img_height_in_cm = self.calculate_size_in_cm(row['Artwork_frame_height'], scale)
        target_img_width_in_cm = self.calculate_size_in_cm(row['Artwork_frame_width'], scale)

        padded_img.save(join(self.temp_dir.name, 'temp_img.jpg'))
        self.document.add_picture(join(self.temp_dir.name, 'temp_img.jpg'), height=Cm(target_img_height_in_cm),
                                  width=Cm(target_img_width_in_cm))

    def _get_padded_img(self, row):
        inp_img = row['Image_file']
        img_real_height = row['Artwork_life_size_height']
        img_real_width = row['Artwork_life_size_width']
        frame_real_height = row['Artwork_frame_height']
        frame_real_width = row['Artwork_frame_width']

        self.image_orientation = ImageOrientation(inp_img)
        picture = self.image_orientation.re_orient()
        width, height = picture.size
        height_padding_ratio = frame_real_height / img_real_height
        width_padding_ratio = frame_real_width / img_real_width

        width_diff = int((width*width_padding_ratio)-width-5)
        height_diff = int((height*height_padding_ratio)-height-5)
        padded_img = ImageOps.expand(picture, border=(width_diff, height_diff, width_diff, height_diff), fill='white')

        return padded_img
